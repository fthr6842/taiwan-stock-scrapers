# 負責證交所新聞
# =============
#Modules
import pandas as pd
import requests
import os
from datetime import datetime
import json
from docx import Document
import time

# Config
save_folder = r"D:\DataPool_D"
today = datetime.today().strftime("%Y%m%d_%H%M%S")
filename = f"{today}_TWSE_Listed_Firm_Announcements.docx"
url = "https://openapi.twse.com.tw/v1/opendata/t187ap04_L"

#Utils
def dev_tools (url:str, encoding: str = "utf-8"):
    """
        爬蟲測試工具
    """
    resp = requests.get(url)
    resp.encoding = encoding
    if resp.status_code == 200:
        print("成功爬取")
    else:
        print("爬取失敗")
    data_list = json.loads(resp.text)
    df = pd.DataFrame(data_list)
    cols = df.columns.tolist()
    # 對調第 5 和第 7 欄
    cols[5], cols[7] = cols[7], cols[5]
    # 重新指定欄位順序
    df = df[cols]
    return df

def doc_process (df, save_folder, filename, title = "TWSE 上市公司重訊"):
    # 建立 Word 文件
    doc = Document()
    # 加入標題
    doc.add_heading(title, level=1)
    # 建立表格 (rows = 資料列數+1, cols = 欄位數)
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    # 設定表格樣式
    table.style = 'Table Grid'
    # 填入欄位名稱
    for j, col_name in enumerate(df.columns):
        table.cell(0, j).text = str(col_name)
    # 填入資料
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            table.cell(i + 1, j).text = str(df.iat[i, j])
    # 存檔
    full_path = os.path.join(save_folder, filename)
    doc.save(full_path)
    return None

# 執行區域
if __name__ == "__main__":
    try:
    
        df = dev_tools(url)
    
        #doc_process (df, save_folder, filename)
        
        # 拆成前面欄位 & 後面兩個欄位
        col_front = df.columns[:-2]
        col_last2 = df.columns[-2:]

        for idx, row in df.iterrows():
            # 先印前面的欄位名稱 & 值
            print(" | ".join(col_front))
            print(" | ".join(row[col_front].astype(str)))

            # 再分行印最後兩個欄位
            for c in col_last2:
                print(f"{c}: {row[c]}")

            print("-" * 80)  # 分隔線（可選）
    
        end = input("按下任意鍵繼續...")
        time.sleep(5)
    except Exception as e: print(f"錯誤原因：{e}\n")